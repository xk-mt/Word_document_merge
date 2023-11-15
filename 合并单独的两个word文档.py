from docx import Document
from docxcompose.composer import Composer

a = Composer(Document(input('[1]word文档：')))
b=Document()
b.add_page_break()
a.append(b)
a.append(Document(input('[2]word文档：')))
a.save('合并结果.docx')
input('合并完成：合并结果.docx\n\n按任意键退出')
