import os
from docx import Document
from docxcompose.composer import Composer

fyf = Document()
fyf.add_page_break()

wjj = input('[当前文件夹直接回车]文件夹：')
wjml = [f1 for f1 in os.listdir(wjj if wjj else None) if os.path.isfile(f1) and f1.endswith('docx') and not f1.startswith('~$')]
hb = Composer(Document(wjml[0]))
wjj = wjj if wjj[-1] in ['/','\\'] else wjj+'/'

for f1 in wjml[1:]:
    hb.append(fyf)
    hb.append(Document(wjj+f1))

hb.save(wjj+"合并结果.docx")
